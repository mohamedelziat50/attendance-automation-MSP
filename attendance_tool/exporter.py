import docx, docx.shared, docx.oxml


class Exporter:
    # Constructor with valid and invalid rows, and the document's title
    def __init__(self, valid_rows, invalid_rows, title="Attendance Report"):
        self.valid_rows = valid_rows
        self.invalid_rows = invalid_rows
        self.title = title

    # Getter for valid_rows
    @property
    def valid_rows(self):
        return self._valid_rows

    # Setter for valid_rows: Empty Lists are allowed
    @valid_rows.setter
    def valid_rows(self, valid_rows):
        if not isinstance(valid_rows, list):
            raise ValueError("Valid rows must be a list")
        # Check if list is not empty and first item is dict
        if valid_rows and not isinstance(valid_rows[0], dict):
            raise ValueError("Valid rows data must be dictionaries")
        self._valid_rows = valid_rows

    # Getter for invalid_rows
    @property
    def invalid_rows(self):
        return self._invalid_rows

    # Setter for invalid_rows: Empty Lists are allowed
    @invalid_rows.setter
    def invalid_rows(self, invalid_rows):
        if not isinstance(invalid_rows, list):
            raise ValueError("Invalid rows must be a list")
        # Check if list is not empty and first item is dict
        if invalid_rows and not isinstance(invalid_rows[0], dict):
            raise ValueError("Invalid rows data must be dictionaries")
        self._invalid_rows = invalid_rows

    # Getter for title
    @property
    def title(self):
        return self._title

    # Setter for title
    @title.setter
    def title(self, title):
        if not isinstance(title, str):
            raise ValueError("Title must be a string")
        if not title.strip():
            raise ValueError("Title cannot be empty")
        self._title = title

    def export_word(self):
        """
        Export the valid and invalid attendance rows to a Word document.

        Returns: str: The file path of the generated Word document

        Raises PermissionError: If the Word document cannot be created or saved
        """
        # Intialize Document
        document = docx.Document()

        # Set page margins to 1 inch on all sides
        for section in document.sections:
            section.top_margin = docx.shared.Inches(1)
            section.bottom_margin = docx.shared.Inches(1)
            section.left_margin = docx.shared.Inches(1)
            section.right_margin = docx.shared.Inches(1)

        # Add and style document heading
        self.__add_document_heading(document)

        # Add spacing after the heading
        document.add_paragraph()

        # Create and setup the attendance table
        table = self.__create_attendance_table(document)

        # Add valid data rows
        self.__add_valid_data_rows(table)

        # Add invalid data rows (highlighted in red)
        self.__add_invalid_data_rows(table)

        # Add error log section if there are invalid rows
        if self.invalid_rows:
            self.__add_error_log(document)

        try:
            # Rename this title later to self.title or smth similiar
            document.save("demo.docx")
        except PermissionError as error:
            # Raised possibly because file is open, and we're trying to save it
            raise PermissionError(error)
        
    def __add_document_heading(self, document):
        """
        Helper function to add and style the document heading.
        
        ARG document: The Word document to add the heading to
        """
        # Add heading
        heading = document.add_heading(
            f"{self.title} - Microsoft Students Partners Club (MSP)",
        )
        
        # Style the heading
        heading_run = heading.runs[0]  # .runs[0] = first text chunk in the heading that we can style
        heading_run.font.name = "Arial"  # Change font family to Arial
        heading_run.font.size = docx.shared.Pt(21)  # Set font size to 21 points
        heading_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # Set text color to black

    def __create_attendance_table(self, document):
        """
        Helper function to create and setup the attendance table with headers.
        
        ARG: document: The Word document to add the table to
            
        Returns: The created and configured table
        """
        # Define table columns
        columns = ["Name", "ID", "Course Code", "Time", "Name of the Doctor"]

        # Create table: 1 header row + len(valid_rows) rows
        table = document.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        
        # Set column widths to prevent text wrapping
        column_widths = [
            docx.shared.Inches(3.0),  # Name column - wider for long names
            docx.shared.Inches(2.0),  # ID column
            docx.shared.Inches(2.0),  # Course Code column
            docx.shared.Inches(2.5),  # Time column
            docx.shared.Inches(3.0)   # Doctor/TA Name column
        ]
        
        for i, width in enumerate(column_widths):
            table.columns[i].width = width
        
        # Set table border color to blue
        self.__set_table_border_color(table)

        # --- Header Row ---
        hdr_cells = table.rows[0].cells
        # Set margins for header cells
        self.__set_cell_margins(hdr_cells)
        
        for i, col_name in enumerate(columns):
            paragraph = hdr_cells[i].paragraphs[0]
            run = paragraph.add_run(col_name)
            run.font.bold = True
            run.font.name = 'Roboto'
            run.font.size = docx.shared.Pt(11.5)
            # paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            
        return table

    def __add_valid_data_rows(self, table):
        """
        Helper function to add valid data rows to the attendance table.
        
        ARG table: The table object to add rows to
        """
        for row in self.valid_rows:
            cells = table.add_row().cells
            cells[0].text = row["Full Name"]
            cells[1].text = row["University ID"]
            cells[2].text = row["Course Code"]
            cells[3].text = row["Course Time"]
            cells[4].text = row["Doctor/TA Name"]
            
            # Set margins for data cells
            self.__set_cell_margins(cells)
            
            # Format valid row text
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = docx.shared.Pt(11.5)
                        run.font.name = 'Roboto'

    def __add_invalid_data_rows(self, table):
        """
        Helper function to add invalid data rows to the attendance table (highlighted in red).
        
        ARG table: The table object to add rows to
        """
        for row in self.invalid_rows:
            cells = table.add_row().cells
            cells[0].text = row["Full Name"]
            cells[1].text = row["University ID"]
            cells[2].text = row["Course Code"]
            cells[3].text = row["Course Time"]
            cells[4].text = row["Doctor/TA Name"]
            
            # Set margins for invalid data cells
            self.__set_cell_margins(cells)
            
            # Highlight invalid row text in red
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red color
                        run.font.size = docx.shared.Pt(11.5)
                        run.font.name = 'Roboto'

    def __set_cell_margins(self, cells):
        """
        Helper function to set cell margins using XML manipulation.
        
        ARG cells: List of table cells to apply margins to
        ARG margin_size: Margin size in dxa units (default 120 = ~6 points, or change up)
        """

        # Specific Margin Sizes
        top_margin = 150
        right_margin = 360
        bottom_margin = 360
        left_margin = 150
        
        for cell in cells:
            cell_element = cell._element
            cell_properties = cell_element.get_or_add_tcPr()
            margins = docx.oxml.parse_xml(f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                        f'<w:top w:w="{top_margin}" w:type="dxa"/>'
                                        f'<w:left w:w="{left_margin}" w:type="dxa"/>'
                                        f'<w:bottom w:w="{bottom_margin}" w:type="dxa"/>'
                                        f'<w:right w:w="{right_margin}" w:type="dxa"/>'
                                        f'</w:tcMar>')
            cell_properties.append(margins)

    def __set_table_border_color(self, table, color="0000FF"):
        """
        Helper function to set table border color using XML manipulation.

        ARG table: The table object to apply border color to
        ARG color: Hex color code (default "0000FF" for blue)
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = docx.oxml.parse_xml(f'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                       f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                                       f'</w:tblBorders>')
        tblPr.append(tblBorders)

    def __add_error_log(self, document):
        """
        Helper function to add validation error log section to the document.
        
        ARG document: The Word document to add the error log to
        """
        # Add some space before the log
        document.add_paragraph()
        
        # Add log heading
        log_heading = document.add_paragraph()
        log_heading_run = log_heading.add_run("Validation Issues Log:")
        log_heading_run.font.bold = True
        log_heading_run.font.size = docx.shared.Pt(14)
        log_heading_run.font.name = 'Roboto'
        log_heading_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red color
        
        # Add each error as a separate paragraph
        for i, row in enumerate(self.invalid_rows, 1):
            if "error" in row and row["error"]:
                error_paragraph = document.add_paragraph()
                
                # Add error number (invalid_row index + 1)
                error_run = error_paragraph.add_run(f"{i}. ")
                error_run.font.bold = True
                error_run.font.name = 'Roboto'
                error_run.font.size = docx.shared.Pt(11)
                
                # Add student identifier (name or ID)
                student_name = row["Full Name"]
                if not student_name.strip():
                    # Default value 'Unknown'
                    if row["University ID"]:
                        student_name = f"Student with ID: {row["University ID"]}"
                    else:
                        student_name = f"Unknown Student"
                
                # Student - 
                name_run = error_paragraph.add_run(f"{student_name} - ")
                name_run.font.bold = True
                name_run.font.name = 'Roboto'
                name_run.font.size = docx.shared.Pt(11)
                
                # Add error message
                error_msg_run = error_paragraph.add_run(row["error"])
                error_msg_run.font.name = 'Roboto'
                error_msg_run.font.size = docx.shared.Pt(11)
                error_msg_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)  # Gray color